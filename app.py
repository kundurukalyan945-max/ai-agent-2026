import streamlit as st
from groq import Groq
from dotenv import load_dotenv
import os
from datetime import datetime
import pandas as pd
import io

load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

st.set_page_config(
    page_title="Kalyan's AI Career Suite",
    page_icon="◆",
    layout="wide",
    initial_sidebar_state="expanded"
)

if "messages" not in st.session_state:
    st.session_state.messages = []
if "score_history" not in st.session_state:
    st.session_state.score_history = []
if "theme" not in st.session_state:
    st.session_state.theme = "dark"
if "last_resume" not in st.session_state:
    st.session_state.last_resume = None
if "job_results" not in st.session_state:
    st.session_state.job_results = None

THEMES = {
    "dark": {"bg": "#0B0E11", "bg_card": "#14181D", "border": "#262C33", "accent": "#FF6A1A", "accent_dim": "#FF6A1A33", "text": "#E8EAED", "text_dim": "#8A919C", "text_faint": "#565D66", "green": "#3DD68C", "btn_text": "#0B0E11"},
    "light": {"bg": "#FAF8F5", "bg_card": "#FFFFFF", "border": "#E4E0D8", "accent": "#D9540C", "accent_dim": "#D9540C22", "text": "#1A1A1A", "text_dim": "#5C5850", "text_faint": "#8C887E", "green": "#1F8C5C", "btn_text": "#FFFFFF"}
}
T = THEMES[st.session_state.theme]

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;500;700&family=Inter:wght@400;500;600;700;800&display=swap');
.stApp {{ background: {T['bg']}; }}
* {{ font-family: 'Inter', sans-serif; }}
#MainMenu, footer, header {{ visibility: hidden; }}
.block-container {{ padding-top: 2rem; max-width: 1100px; }}
.suite-header {{ display: flex; align-items: baseline; gap: 14px; margin-bottom: 2px; }}
.suite-mark {{ color: {T['accent']}; font-family: 'JetBrains Mono', monospace; font-size: 28px; font-weight: 700; }}
.suite-title {{ font-size: 30px; font-weight: 800; color: {T['text']}; letter-spacing: -0.5px; }}
.suite-sub {{ font-family: 'JetBrains Mono', monospace; font-size: 13px; color: {T['text_faint']}; margin-bottom: 28px; margin-top: 4px; }}
.suite-sub .dot {{ color: {T['green']}; }}
.status-strip {{ display: flex; border: 1px solid {T['border']}; border-radius: 10px; overflow: hidden; background: {T['bg_card']}; }}
.status-cell {{ flex: 1; padding: 14px 18px; }}
.status-label {{ font-family: 'JetBrains Mono', monospace; font-size: 10px; color: {T['text_faint']}; text-transform: uppercase; letter-spacing: 1px; margin-bottom: 4px; }}
.status-value {{ font-size: 14px; color: {T['text']}; font-weight: 600; }}
.status-value.accent {{ color: {T['accent']}; }}
.panel-head {{ display: flex; align-items: center; gap: 10px; margin-bottom: 4px; }}
.panel-tag {{ font-family: 'JetBrains Mono', monospace; font-size: 11px; color: {T['accent']}; background: {T['accent_dim']}; padding: 3px 9px; border-radius: 5px; font-weight: 700; }}
.panel-title {{ font-size: 20px; font-weight: 700; color: {T['text']}; }}
.panel-desc {{ color: {T['text_dim']}; font-size: 13.5px; margin: 4px 0 20px 0; }}
.stTextInput input, .stTextArea textarea, .stSelectbox > div > div {{ background: {T['bg_card']} !important; border: 1px solid {T['border']} !important; color: {T['text']} !important; border-radius: 8px !important; }}
.stButton button {{ background: {T['accent']} !important; color: {T['btn_text']} !important; border: none !important; border-radius: 8px !important; font-weight: 700 !important; padding: 10px 22px !important; }}
.stButton button:hover {{ filter: brightness(1.15); }}
.stDownloadButton button {{ background: transparent !important; color: {T['accent']} !important; border: 1px solid {T['accent']} !important; border-radius: 8px !important; }}
.stTabs [data-baseweb="tab-list"] {{ gap: 4px; background: {T['bg_card']}; padding: 5px; border-radius: 10px; border: 1px solid {T['border']}; flex-wrap: wrap; }}
.stTabs [data-baseweb="tab"] {{ background: transparent; border-radius: 7px; color: {T['text_dim']}; font-weight: 600; padding: 10px 16px; }}
.stTabs [aria-selected="true"] {{ background: {T['accent']} !important; color: {T['btn_text']} !important; }}
.output-card {{ background: {T['bg_card']}; border: 1px solid {T['border']}; border-radius: 10px; padding: 22px 24px; margin-top: 18px; color: {T['text']}; }}
.output-card-label {{ font-family: 'JetBrains Mono', monospace; font-size: 11px; color: {T['green']}; text-transform: uppercase; margin-bottom: 12px; }}
.score-badge {{ display: inline-flex; font-family: 'JetBrains Mono', monospace; font-size: 22px; font-weight: 700; padding: 8px 18px; border-radius: 8px; margin-bottom: 14px; }}
.score-high {{ background: #3DD68C22; color: #3DD68C; border: 1px solid #3DD68C44; }}
.score-mid {{ background: #FFB13C22; color: #FFB13C; border: 1px solid #FFB13C44; }}
.score-low {{ background: #FF5C5C22; color: #FF5C5C; border: 1px solid #FF5C5C44; }}
[data-testid="stSidebar"] {{ background: {T['bg_card']}; border-right: 1px solid {T['border']}; }}
[data-testid="stChatMessage"] {{ background: {T['bg_card']} !important; border: 1px solid {T['border']} !important; border-radius: 10px !important; }}
hr {{ border-color: {T['border']} !important; }}
</style>
""", unsafe_allow_html=True)

with st.sidebar:
    st.markdown(f"<div style=\"font-family:'JetBrains Mono',monospace;color:{T['accent']};font-size:13px;font-weight:700;\">◆ KALYAN.AI</div>", unsafe_allow_html=True)
    theme_choice = st.radio("Theme", ["dark", "light"], index=0 if st.session_state.theme == "dark" else 1, horizontal=True)
    if theme_choice != st.session_state.theme:
        st.session_state.theme = theme_choice
        st.rerun()
    st.markdown("---")
    st.markdown("**Builder**")
    st.markdown("Kundura Kalyan")
    st.markdown("B.Tech CSE (AI/ML) · SVCE")
    st.markdown("---")
    st.markdown("**Stack**")
    st.markdown("`Python` `Groq` `LLaMA 3.3`")
    if st.session_state.score_history:
        st.markdown("---")
        avg = sum(s["score"] for s in st.session_state.score_history) / len(st.session_state.score_history)
        st.markdown("**Interview avg**")
        st.markdown(f"<span style='font-family:JetBrains Mono;font-size:24px;color:{T['accent']};font-weight:700;'>{avg:.1f}/10</span>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("[GitHub Repo →](https://github.com/kundurukalyan945-max/ai-agent-2026)")

st.markdown(f"""
<div class="suite-header"><span class="suite-mark">◆</span><span class="suite-title">AI Career Suite</span></div>
<div class="suite-sub"><span class="dot">●</span> ONLINE &nbsp;·&nbsp; built by Kalyan &nbsp;·&nbsp; powered by LLaMA 3.3-70B</div>
""", unsafe_allow_html=True)

c1, c2, c3, c4 = st.columns(4)
cells = [("Agents", "05 active", True), ("Model", "LLaMA 3.3 70B", False), ("Target", "AI/ML · Bengaluru", False), ("Cost", "₹0", True)]
for col, (label, value, accent) in zip([c1, c2, c3, c4], cells):
    cls = "accent" if accent else ""
    col.markdown(f'<div class="status-strip"><div class="status-cell"><div class="status-label">{label}</div><div class="status-value {cls}">{value}</div></div></div>', unsafe_allow_html=True)

st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
tab1, tab2, tab3, tab4, tab5 = st.tabs(["COACH", "RESUME", "INTERVIEW", "JOBS", "CREW"])
# ══════════════════════════════════════════════════════════
# TAB 1 — COACH
# ══════════════════════════════════════════════════════════
with tab1:
    st.markdown('<div class="panel-head"><span class="panel-tag">AGENT 01</span><span class="panel-title">Career Coach</span></div><div class="panel-desc">Persistent memory · Upload a PDF to discuss it.</div>', unsafe_allow_html=True)

    uploaded_pdf = st.file_uploader("Upload a PDF (resume, job posting, anything)", type=["pdf"], key="coach_pdf")
    pdf_text = ""
    if uploaded_pdf:
        try:
            from pypdf import PdfReader
            reader = PdfReader(uploaded_pdf)
            pdf_text = "\n".join([page.extract_text() or "" for page in reader.pages])
            st.success(f"📄 {uploaded_pdf.name} loaded — {len(reader.pages)} pages. Ask about it below.")
        except Exception as e:
            st.error(f"Couldn't read PDF: {e}")

    st.markdown('<div class="chat-scroll">', unsafe_allow_html=True)
    for msg in st.session_state.messages:
        bubble_class = "bubble-user" if msg["role"] == "user" else "bubble-ai"
        align = "flex-end" if msg["role"] == "user" else "flex-start"
        st.markdown(f'<div style="display:flex; justify-content:{align}; margin:6px 0;"><div class="{bubble_class}">{msg["content"]}</div></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if prompt := st.chat_input("Ask your career coach anything..."):
        full_prompt = prompt
        if pdf_text:
            full_prompt = f"[Context from uploaded PDF '{uploaded_pdf.name}']:\n{pdf_text[:4000]}\n\n[Question]: {prompt}"

        st.session_state.messages.append({"role": "user", "content": prompt})

        with st.spinner("Thinking..."):
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "system", "content": "You are Kalyan's AI career coach. He is a B.Tech CSE AI/ML student in Bengaluru targeting AI/ML jobs in 2026. Be encouraging, direct, and practical. If PDF context is given, use it to answer."}] + 
                [{"role": m["role"], "content": full_prompt if i == len(st.session_state.messages)-1 and m["role"]=="user" else m["content"]} for i, m in enumerate(st.session_state.messages)]
            )
            reply = response.choices[0].message.content
            st.session_state.messages.append({"role": "assistant", "content": reply})
        st.rerun()

# ══════════════════════════════════════════════════════════
# TAB 2 — RESUME
# ══════════════════════════════════════════════════════════
with tab2:
    st.markdown('<div class="panel-head"><span class="panel-tag">AGENT 02</span><span class="panel-title">Resume Tailor</span></div><div class="panel-desc">Paste a job description. Download as .txt or .docx.</div>', unsafe_allow_html=True)

    job_desc = st.text_area("Job description", height=160, label_visibility="collapsed", placeholder="Paste the full job description here...")

    if st.button("Tailor my resume", key="tailor_btn"):
        if job_desc:
            with st.spinner("Analyzing requirements..."):
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {"role": "system", "content": "You are an expert ATS resume writer."},
                        {"role": "user", "content": f"""My resume: B.Tech CSE AI/ML, SVCE Bengaluru 2026, CGPA 7.3.
Project: Deep Learning Object Detection for Visually Impaired (YOLOv8, TTS, Streamlit, 93-95% accuracy).
Skills: Python, ML, Deep Learning, LangChain, Groq API, AI Agents.

Job description: {job_desc}

Give me:
1. Top 5 ATS keywords to include
2. 3 rewritten project bullet points matching this job
3. A 3-line professional summary for this role"""}
                    ]
                )
                st.session_state.last_resume = response.choices[0].message.content
        else:
            st.warning("Paste a job description first.")

    if st.session_state.last_resume:
        st.markdown(f'<div class="output-card"><div class="output-card-label">● output ready</div>{st.session_state.last_resume}</div>', unsafe_allow_html=True)
        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button("↓ Download as .txt", data=st.session_state.last_resume, file_name=f"tailored_resume_{datetime.now().strftime('%Y%m%d')}.txt", mime="text/plain")
        with dl2:
            try:
                from docx import Document
                doc = Document()
                doc.add_heading("Tailored Resume Content", level=1)
                for line in st.session_state.last_resume.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                st.download_button("↓ Download as .docx", data=buf, file_name=f"tailored_resume_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except ImportError:
                st.caption("Run: pip install python-docx")
                # ══════════════════════════════════════════════════════════
# TAB 3 — INTERVIEW (with progress tracker)
# ══════════════════════════════════════════════════════════
with tab3:
    st.markdown('<div class="panel-head"><span class="panel-tag">AGENT 03</span><span class="panel-title">Interview Prep</span></div><div class="panel-desc">Practice with real questions. Scores tracked across your session.</div>', unsafe_allow_html=True)

    questions = [
        "Tell me about yourself and why you want to work in AI/ML?",
        "What is machine learning in simple terms?",
        "What is the difference between supervised and unsupervised learning?",
        "Tell me about your Object Detection project and challenges you faced.",
        "What is overfitting and how do you fix it?"
    ]

    question = st.selectbox("Question", questions, label_visibility="collapsed")
    answer = st.text_area("Answer", height=120, label_visibility="collapsed", placeholder="Type your answer here...")

    if st.button("Score my answer", key="score_btn"):
        if answer:
            with st.spinner("Evaluating..."):
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[
                        {"role": "system", "content": "You are a strict technical interviewer. Score out of 10. Be direct and fast."},
                        {"role": "user", "content": f"Question: {question}\nAnswer: {answer}\n\nRespond with:\nSCORE: X/10\nGOOD: (one line)\nIMPROVE: (one line)\nBETTER ANSWER: (2-3 lines)"}
                    ]
                )
                feedback = response.choices[0].message.content
                lines = feedback.split('\n')
                score_line = [l for l in lines if 'SCORE' in l]
                if score_line:
                    score = int(score_line[0].split('/')[0].replace('SCORE:', '').strip())
                    st.session_state.score_history.append({"time": datetime.now().strftime("%H:%M:%S"), "question": question[:30] + "...", "score": score})
                    css_class = "score-high" if score >= 7 else "score-mid" if score >= 5 else "score-low"
                    st.markdown(f'<div class="score-badge {css_class}">{score_line[0].strip()}</div>', unsafe_allow_html=True)
                    feedback_clean = '\n'.join([l for l in lines if 'SCORE' not in l])
                    st.markdown(f'<div class="output-card">{feedback_clean}</div>', unsafe_allow_html=True)
        else:
            st.warning("Type your answer first.")

    if st.session_state.score_history:
        st.markdown("<div style='height:24px'></div>", unsafe_allow_html=True)
        st.markdown('<div class="panel-head"><span class="panel-tag">PROGRESS</span><span class="panel-title" style="font-size:16px;">Score History</span></div>', unsafe_allow_html=True)
        df = pd.DataFrame(st.session_state.score_history)
        st.line_chart(df.set_index("time")["score"], height=200)
        st.dataframe(df, use_container_width=True, hide_index=True)

# ══════════════════════════════════════════════════════════
# TAB 4 — JOB SCRAPER
# ══════════════════════════════════════════════════════════
with tab4:
    st.markdown('<div class="panel-head"><span class="panel-tag">AGENT 04</span><span class="panel-title">Job Scraper</span></div><div class="panel-desc">Search live AI/ML job postings in Bengaluru.</div>', unsafe_allow_html=True)

    jc1, jc2 = st.columns([3, 1])
    with jc1:
        search_term = st.text_input("Search term", value="AI ML Python developer", label_visibility="collapsed")
    with jc2:
        search_btn = st.button("Search jobs", key="job_search_btn")

    if search_btn:
        with st.spinner("Searching Indeed... (~30 sec)"):
            try:
                from jobspy import scrape_jobs
                jobs = scrape_jobs(site_name=["indeed"], search_term=search_term, location="Bengaluru, India", results_wanted=15, hours_old=72, country_indeed="India")
                if not jobs.empty:
                    st.session_state.job_results = jobs[["title", "company", "location", "date_posted", "job_url"]].copy()
                else:
                    st.warning("No jobs found. Try a different search term.")
            except ImportError:
                st.error("Run: pip install python-jobspy")
            except Exception as e:
                st.error(f"Search failed: {e}")

    if st.session_state.job_results is not None:
        st.markdown(f'<div class="output-card-label" style="margin-top:18px;">● {len(st.session_state.job_results)} jobs found</div>', unsafe_allow_html=True)
        st.dataframe(st.session_state.job_results, use_container_width=True, hide_index=True, column_config={"job_url": st.column_config.LinkColumn("Apply Link")})
        csv = st.session_state.job_results.to_csv(index=False).encode('utf-8')
        st.download_button("↓ Download as CSV", data=csv, file_name=f"jobs_{datetime.now().strftime('%Y%m%d')}.csv", mime="text/csv")

# ══════════════════════════════════════════════════════════
# TAB 5 — MULTI-AGENT CREW
# ══════════════════════════════════════════════════════════
with tab5:
    st.markdown('<div class="panel-head"><span class="panel-tag">AGENT 05</span><span class="panel-title">Multi-Agent Crew</span></div><div class="panel-desc">3 agents work in sequence: Researcher → Resume Writer → Interview Coach.</div>', unsafe_allow_html=True)

    cc1, cc2 = st.columns(2)
    with cc1:
        crew_company = st.text_input("Company", placeholder="e.g. Google")
    with cc2:
        crew_role = st.text_input("Role", placeholder="e.g. AI Engineer")

    if st.button("Launch crew", key="crew_btn"):
        if crew_company and crew_role:
            progress = st.progress(0, text="Agent 1 — Researching company...")
            research = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "system", "content": "You are an expert job market researcher."}, {"role": "user", "content": f"Research {crew_company} for someone applying for {crew_role}. Cover: what they do, tech stack, culture, what they look for in candidates, top 5 resume keywords."}]).choices[0].message.content
            progress.progress(33, text="Agent 2 — Tailoring resume...")
            resume = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "system", "content": "You are an ATS resume expert."}, {"role": "user", "content": f"Based on this research:\n{research}\n\nTailor Kalyan's resume for {crew_role} at {crew_company}. Kalyan: B.Tech CSE AI/ML, AI Career Suite builder, YOLOv8 project (93-95% accuracy), Python/ML/LangChain/Groq skills, SmartBridge/JPMorgan/Deloitte internships. Give: 3-line summary, 5 bullet points, 8 ATS keywords."}]).choices[0].message.content
            progress.progress(66, text="Agent 3 — Preparing interview...")
            interview = client.chat.completions.create(model="llama-3.3-70b-versatile", messages=[{"role": "system", "content": "You are a senior technical interview coach."}, {"role": "user", "content": f"Using this research:\n{research}\n\nAnd resume:\n{resume}\n\nPrepare Kalyan for {crew_role} interview at {crew_company}: top 5 technical questions with answers, top 3 behavioral questions with STAR answers, 3 smart questions to ask, one insider tip."}]).choices[0].message.content
            progress.progress(100, text="Done!")

            st.markdown(f'<div class="output-card"><div class="output-card-label">● agent 1 — research</div>{research}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="output-card"><div class="output-card-label">● agent 2 — tailored resume</div>{resume}</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="output-card"><div class="output-card-label">● agent 3 — interview prep</div>{interview}</div>', unsafe_allow_html=True)

            full_report = f"CREW REPORT: {crew_role} at {crew_company}\n\n=== RESEARCH ===\n{research}\n\n=== RESUME ===\n{resume}\n\n=== INTERVIEW PREP ===\n{interview}"
            st.download_button("↓ Download full report", data=full_report, file_name=f"crew_{crew_company}_{crew_role}.txt".replace(" ", "_"), mime="text/plain")
        else:
            st.warning("Enter both company and role.")